# src/extractor.py
import pdfplumber
import docx
import filetype
import re
import nltk
from pathlib import Path
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTChar
from nltk.sentiment.vader import SentimentIntensityAnalyzer
from langdetect import detect

# Initialize NLP
nltk.download('vader_lexicon', quiet=True)

class ResearchExtractor:
    """Base extractor with text, structure, sentiment, and entities."""
    def __init__(self, file_path):
        self.path = Path(file_path)
        self.analyzer = SentimentIntensityAnalyzer()

    def _sanitize_text(self, text):
        if not text: return ""
        return "".join(c for c in text if c.isprintable() or c in "\n\r\t")

    def process(self):
        kind = filetype.guess(str(self.path))
        detected_ext = kind.extension if kind else self.path.suffix.replace(".", "").lower()
        mime = kind.mime if kind else "application/octet-stream"

        results = {
            "text": "",
            "structure": [],
            "entities": {},
            "sentiment": {},
            "tables_found": 0,
            "status": "success",
            "metadata": {
                "filename": self.path.name,
                "detected_format": detected_ext.upper(),
                "mime": mime,
                "integrity": "✅ Match" if self.path.suffix.lower() == f".{detected_ext}" else "⚠️ Mismatch",
                "language": "unknown"
            }
        }

        try:
            if detected_ext == 'pdf':
                results["text"], results["tables_found"] = self._parse_pdf_text()
                results["structure"] = self._get_pdf_structure()
            elif detected_ext in ['docx', 'doc']:
                results["text"], results["structure"] = self._get_docx_structure()
            else:
                results["text"] = self.path.read_text(errors='ignore')

            results["text"] = self._sanitize_text(results["text"])
            results["sentiment"] = self._analyze_sentiment(results["text"])
            results["entities"] = self._discover_entities(results["text"])

            # Detect language
            try:
                results["metadata"]["language"] = detect(results["text"])
            except: results["metadata"]["language"] = "unknown"

        except Exception as e:
            results["status"] = f"Error: {str(e)}"

        return results

    def _get_pdf_structure(self):
        structure = []
        try:
            for page_layout in extract_pages(self.path):
                for element in page_layout:
                    if isinstance(element, LTTextContainer):
                        text = element.get_text().strip()
                        if text:
                            is_bold = any("bold" in str(line).lower() for line in element)
                            structure.append({"text": text, "is_bold": is_bold, "level": 1 if is_bold else 2})
        except: pass
        return structure

    def _get_docx_structure(self):
        doc = docx.Document(self.path)
        full_text = [p.text for p in doc.paragraphs]
        structure = [{"text": p.text, "is_heading": p.style.name.startswith('Heading'), "level": int(p.style.name[-1]) if p.style.name.startswith('Heading') else 2} 
                     for p in doc.paragraphs if p.text.strip()]
        return "\n".join(full_text), structure

    def _analyze_sentiment(self, text):
        if not text.strip(): return {"compound": 0, "label": "neutral"}
        s = self.analyzer.polarity_scores(text)
        label = "positive" if s['compound'] >= 0.05 else "negative" if s['compound'] <= -0.05 else "neutral"
        return {**s, "label": label}

    def _discover_entities(self, text):
        return {
            "currency": list(set(re.findall(r'\$\s?\d+(?:,\d{3})*(?:\.\d{2})?', text))),
            "project_ids": list(set(re.findall(r'REF-\d{4}', text))),
            "emails": list(set(re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)))
        }

    def _parse_pdf_text(self):
        text_content = ""
        t_count = 0
        with pdfplumber.open(self.path) as pdf:
            for page in pdf.pages:
                text_content += (page.extract_text() or "") + "\n"
                t_count += len(page.extract_tables())
        return text_content, t_count


# ================================
# Industrial-Level Advanced Extractor
# ================================
class AdvancedResearchExtractor(ResearchExtractor):
    """Extends ResearchExtractor with table extraction for PDF and DOCX."""

    def _parse_pdf_tables(self):
        """Extracts all tables from a PDF and returns a list of dicts."""
        tables_list = []
        try:
            with pdfplumber.open(self.path) as pdf:
                for page_number, page in enumerate(pdf.pages, 1):
                    page_tables = page.extract_tables()
                    for tbl in page_tables:
                        tables_list.append({
                            "page": page_number,
                            "table": tbl
                        })
        except Exception as e:
            print(f"Error extracting PDF tables: {e}")
        return tables_list

    def _parse_docx_tables(self):
        """Extracts all tables from a DOCX file and returns a list of dicts."""
        tables_list = []
        try:
            doc = docx.Document(self.path)
            for idx, table in enumerate(doc.tables, 1):
                tbl_data = []
                for row in table.rows:
                    tbl_data.append([cell.text.strip() for cell in row.cells])
                tables_list.append({
                    "page": idx,  # DOCX doesn't have pages, use table index
                    "table": tbl_data
                })
        except Exception as e:
            print(f"Error extracting DOCX tables: {e}")
        return tables_list
