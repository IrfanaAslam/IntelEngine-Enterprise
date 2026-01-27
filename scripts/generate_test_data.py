import os
import random
import pandas as pd
from pathlib import Path
from faker import Faker
from docx import Document

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

# Initialization
fake = Faker()
styles = getSampleStyleSheet()

# 1. SETUP DIRECTORIES
BASE_DIR = Path.cwd()
RAW_DATA_DIR = BASE_DIR / "data" / "raw"
RAW_DATA_DIR.mkdir(parents=True, exist_ok=True)

def create_pdf(filepath):
    """Generates a structured PDF with a title, body text, and a data table."""
    doc = SimpleDocTemplate(str(filepath), pagesize=letter)
    elements = []

    # Title
    elements.append(Paragraph(f"Research Analysis: {fake.bs().title()}", styles['Title']))
    elements.append(Spacer(1, 12))

    # Narrative Text
    elements.append(Paragraph(fake.paragraph(nb_sentences=5), styles['Normal']))
    elements.append(Spacer(1, 12))

    # Structured Table
    data = [["Feature ID", "Observation", "Metric Value"]]
    for _ in range(5):
        data.append([
            f"REF-{random.randint(1000, 9999)}",
            fake.word().capitalize(),
            f"{random.uniform(0.1, 99.9):.2f}"
        ])
    
    table = Table(data, colWidths=[100, 200, 100])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.cadetblue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(table)
    
    doc.build(elements)

def create_docx(filepath):
    """Generates a Microsoft Word document."""
    doc = Document()
    doc.add_heading(f"Study: {fake.catch_phrase()}", 0)
    p = doc.add_paragraph(fake.text(max_nb_chars=1200))
    p.bold = True
    doc.add_page_break()
    doc.save(filepath)

def create_csv(filepath):
    """Generates a CSV file with random data."""
    df = pd.DataFrame({
        "Timestamp": [fake.date_time_this_year() for _ in range(20)],
        "Sample_Size": [random.randint(50, 500) for _ in range(20)],
        "Coefficient": [random.random() for _ in range(20)],
        "Status": [random.choice(["Control", "Test", "Placebo"]) for _ in range(20)]
    })
    df.to_csv(filepath, index=False)

def main(total_files=1000):
    print(f"📁 Target Folder: {RAW_DATA_DIR.absolute()}")
    print(f"🚀 Starting generation of {total_files} files...")

    counts = {"pdf": 0, "docx": 0, "csv": 0}
    
    for i in range(total_files):
        # Determine file type distribution
        # 60% PDFs, 20% Word, 20% CSV
        choice = random.choices(['pdf', 'docx', 'csv'], weights=[60, 20, 20])[0]
        filename = f"document_{i:04d}.{choice}"
        filepath = RAW_DATA_DIR / filename

        try:
            if choice == 'pdf':
                create_pdf(filepath)
            elif choice == 'docx':
                create_docx(filepath)
            elif choice == 'csv':
                create_csv(filepath)
            
            counts[choice] += 1
        except Exception as e:
            print(f"❌ Error creating {filename}: {e}")

        # Progress indicator every 100 files
        if (i + 1) % 100 == 0:
            print(f"Progress: {i + 1}/{total_files} files generated...")

    print("\n--- GENERATION SUMMARY ---")
    print(f"PDFs Created:  {counts['pdf']}")
    print(f"DOCX Created: {counts['docx']}")
    print(f"CSVs Created:  {counts['csv']}")
    print(f"Total:         {sum(counts.values())}")
    print(f"Location:      {RAW_DATA_DIR}")

if __name__ == "__main__":
    main(1000)